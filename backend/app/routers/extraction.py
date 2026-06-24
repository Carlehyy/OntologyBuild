"""
Document Extraction Router

Handles document upload, LLM-based extraction, and human review workflow.
This is the entry point for the "Document -> Ontology" pipeline.
"""

import os
import uuid
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session

from app.database import get_db
from app.models import (
    Document, ExtractionResult, ObjectType, RelationType, Entity, Relation,
    ExtractionStatus, AuditLog, FeedbackRecord
)
from app.schemas import (
    DocumentCreate, DocumentOut,
    ExtractionResultOut, ExtractionReview,
    FeedbackCreate,
)
from app.services.llm_service import get_llm_service
from app.services.graph_service import get_graph_service
from app.config import get_settings

router = APIRouter(prefix="/extraction", tags=["Extraction"])
settings = get_settings()


@router.get("/domain/{domain_id}/documents", response_model=List[DocumentOut])
def list_documents(
    domain_id: str,
    db: Session = Depends(get_db),
    status: Optional[str] = None,
):
    """List documents for a domain."""
    query = db.query(Document).filter(Document.domain_id == domain_id)
    if status:
        query = query.filter(Document.status == status)
    return query.order_by(Document.created_at.desc()).all()


@router.post("/domain/{domain_id}/documents", response_model=DocumentOut)
async def upload_document(
    domain_id: str,
    file: UploadFile = File(...),
    extraction_config: Optional[str] = "{}",
    db: Session = Depends(get_db),
):
    """Upload a document for extraction."""
    import json

    # Save file
    file_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename)[1]
    stored_filename = f"{file_id}{ext}"
    file_path = os.path.join("./data/uploads", stored_filename)

    os.makedirs("./data/uploads", exist_ok=True)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # Read text content (basic handling)
    text_content = ""
    try:
        if ext.lower() in [".txt", ".md", ".csv"]:
            text_content = content.decode("utf-8", errors="ignore")
        elif ext.lower() in [".pdf"]:
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_content = "\n".join([page.get_text() for page in doc])
            except ImportError:
                text_content = "[PDF parsing requires PyMuPDF]"
        elif ext.lower() in [".docx", ".doc"]:
            try:
                import docx
                doc = docx.Document(file_path)
                text_content = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                text_content = "[DOCX parsing requires python-docx]"
        else:
            text_content = content.decode("utf-8", errors="ignore")[:50000]
    except Exception as e:
        text_content = f"[Error reading content: {str(e)}]"

    # Create document record
    config = {}
    try:
        config = json.loads(extraction_config) if extraction_config else {}
    except json.JSONDecodeError:
        pass

    doc = Document(
        domain_id=domain_id,
        filename=stored_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=len(content),
        mime_type=file.content_type,
        content_text=text_content,
        extraction_config=config,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Audit log
    db.add(AuditLog(
        action="upload",
        resource_type="document",
        resource_id=doc.id,
        domain_id=domain_id,
        details={"filename": file.filename, "size": len(content)},
    ))
    db.commit()

    return doc


@router.post("/documents/{document_id}/extract")
async def run_extraction(document_id: str, db: Session = Depends(get_db)):
    """Run LLM extraction on a document."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if not doc.content_text or len(doc.content_text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Document has no extractable text")

    # Update status
    doc.status = ExtractionStatus.PROCESSING
    db.commit()

    try:
        # Get ontology for this domain
        object_types = db.query(ObjectType).filter(
            ObjectType.domain_id == doc.domain_id,
            ObjectType.is_active == True,
        ).all()

        relation_types = db.query(RelationType).filter(
            RelationType.domain_id == doc.domain_id,
            RelationType.is_active == True,
        ).all()

        # Format for LLM
        ot_list = []
        for ot in object_types:
            props = [{"name": p.name, "data_type": p.data_type}
                     for p in ot.properties]
            ot_list.append({
                "name": ot.name,
                "description": ot.description or "",
                "properties": props,
            })

        rt_list = []
        for rt in relation_types:
            rt_list.append({
                "name": rt.name,
                "description": rt.description or "",
                "source_type": rt.source_type.name if rt.source_type else "",
                "target_type": rt.target_type.name if rt.target_type else "",
            })

        # Run LLM extraction
        llm = get_llm_service()

        # Truncate text if too long
        text = doc.content_text[:15000]  # Limit to ~15k chars

        candidates = await llm.extract_from_text(text, ot_list, rt_list)

        # Create extraction results
        entity_count = 0
        relation_count = 0

        for candidate in candidates:
            # Find matching object type
            ot_id = None
            if candidate.object_type:
                for ot in object_types:
                    if ot.name.lower() == candidate.object_type.lower():
                        ot_id = ot.id
                        break

            # Find matching relation type
            rt_id = None
            if candidate.relation_type:
                for rt in relation_types:
                    if rt.name.lower() == candidate.relation_type.lower():
                        rt_id = rt.id
                        break

            result = ExtractionResult(
                document_id=doc.id,
                result_type=candidate.candidate_type,
                candidate_object_type_id=ot_id,
                candidate_name=candidate.name,
                candidate_properties=candidate.properties,
                candidate_relation_type_id=rt_id,
                candidate_source_name=candidate.source_name,
                candidate_target_name=candidate.target_name,
                confidence=candidate.confidence,
                llm_reasoning=candidate.reasoning,
                status="pending",
            )
            db.add(result)

            if candidate.candidate_type == "entity":
                entity_count += 1
            else:
                relation_count += 1

        # Update document
        doc.status = ExtractionStatus.COMPLETED
        doc.extracted_entities_count = entity_count
        doc.extracted_relations_count = relation_count
        db.commit()

        return {
            "success": True,
            "message": f"Extraction complete: {entity_count} entities, {relation_count} relations",
            "entity_count": entity_count,
            "relation_count": relation_count,
        }

    except Exception as e:
        doc.status = ExtractionStatus.FAILED
        db.commit()
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")


@router.get("/documents/{document_id}/results", response_model=List[ExtractionResultOut])
def get_extraction_results(
    document_id: str,
    db: Session = Depends(get_db),
    status: Optional[str] = None,
):
    """Get extraction results for a document."""
    query = db.query(ExtractionResult).filter(ExtractionResult.document_id == document_id)
    if status:
        query = query.filter(ExtractionResult.status == status)
    return query.all()


@router.post("/results/review")
def review_extraction_result(data: ExtractionReview, db: Session = Depends(get_db)):
    """
    Review an extraction result (human-in-the-loop).
    Approved entities get promoted to the graph.
    """
    result = db.query(ExtractionResult).filter(ExtractionResult.id == data.result_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="Extraction result not found")

    result.status = data.action
    result.review_action = data.action
    result.review_comment = data.comment

    doc = result.document
    graph_svc = get_graph_service()

    if data.action == "approved":
        if result.result_type == "entity" and result.candidate_name:
            # Create entity in graph
            entity = Entity(
                domain_id=doc.domain_id,
                object_type_id=result.candidate_object_type_id,
                name=result.candidate_name,
                properties=result.candidate_properties or {},
                source_document_id=doc.id,
                confidence=result.confidence or 0.5,
                is_verified=True,
            )
            db.add(entity)
            db.flush()

            # Sync to graph
            ot = db.query(ObjectType).filter(ObjectType.id == result.candidate_object_type_id).first()
            graph_svc.sync_entity(
                domain_id=doc.domain_id,
                entity_id=entity.id,
                object_type_id=result.candidate_object_type_id or "",
                object_type_name=ot.name if ot else "Unknown",
                name=result.candidate_name,
                properties=result.candidate_properties or {},
                confidence=result.confidence or 0.5,
                is_verified=True,
            )

            # Record feedback
            db.add(FeedbackRecord(
                domain_id=doc.domain_id,
                feedback_type="extraction",
                target_id=result.id,
                target_type="extraction_result",
                verdict="useful",
                context={"entity_id": entity.id, "entity_name": entity.name},
            ))

        elif result.result_type == "relation":
            # Find or create source/target entities
            # This is simplified - full implementation would do entity resolution
            pass

    elif data.action == "rejected":
        # Record negative feedback
        db.add(FeedbackRecord(
            domain_id=doc.domain_id,
            feedback_type="extraction",
            target_id=result.id,
            target_type="extraction_result",
            verdict="false_positive",
            correction_data=data.modifications,
            notes=data.comment,
        ))

    elif data.action == "modified" and data.modifications:
        # Apply modifications and create entity
        result.candidate_name = data.modifications.get("name", result.candidate_name)
        result.candidate_properties = data.modifications.get("properties", result.candidate_properties)
        if result.candidate_object_type_id is None and data.modifications.get("object_type_id"):
            result.candidate_object_type_id = data.modifications["object_type_id"]

        # Now create the entity with modified data
        if result.result_type == "entity" and result.candidate_name:
            entity = Entity(
                domain_id=doc.domain_id,
                object_type_id=result.candidate_object_type_id,
                name=result.candidate_name,
                properties=result.candidate_properties or {},
                source_document_id=doc.id,
                confidence=(result.confidence or 0.5) * 0.9,  # Slightly lower for modified
                is_verified=True,
            )
            db.add(entity)
            db.flush()

            ot = db.query(ObjectType).filter(ObjectType.id == result.candidate_object_type_id).first()
            graph_svc.sync_entity(
                domain_id=doc.domain_id,
                entity_id=entity.id,
                object_type_id=result.candidate_object_type_id or "",
                object_type_name=ot.name if ot else "Unknown",
                name=result.candidate_name,
                properties=result.candidate_properties or {},
                confidence=(result.confidence or 0.5) * 0.9,
                is_verified=True,
            )

            db.add(FeedbackRecord(
                domain_id=doc.domain_id,
                feedback_type="extraction",
                target_id=result.id,
                target_type="extraction_result",
                verdict="needs_correction",
                correction_data=data.modifications,
                context={"entity_id": entity.id},
            ))

    db.commit()
    return {"success": True, "message": f"Result {data.action}"}


@router.post("/documents/{document_id}/review")
def mark_document_reviewed(document_id: str, db: Session = Depends(get_db)):
    """Mark a document as fully reviewed."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    doc.status = ExtractionStatus.REVIEWED
    db.commit()

    return {"success": True, "message": "Document marked as reviewed"}


@router.delete("/documents/{document_id}")
def delete_document(document_id: str, db: Session = Depends(get_db)):
    """Delete a document and its extraction results."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Delete file
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    db.delete(doc)
    db.commit()
    return {"success": True, "message": "Document deleted"}
