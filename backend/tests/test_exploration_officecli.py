"""业务探索 OfficeCLI 适配器的安全边界与版本语义。"""
from __future__ import annotations

import io
import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import HTTPException

from app.config import settings
from app.exploration import officecli as O
from app.exploration import workspace as W
from app.exploration.models import ExplorationAttachment, ExplorationSession
from app.exploration.toolkit import OFFICE_TOOL, ExplorationToolRunner

BASE = "/api/v2/exploration"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture
def exploration_session(client, auth_headers, db):
    response = client.post(f"{BASE}/sessions", headers=auth_headers, json={})
    assert response.status_code == 201, response.text
    session_id = response.json()["data"]["id"]
    return db.query(ExplorationSession).filter_by(id=session_id).one()


def _docx_bytes(text: str = "原始条款") -> bytes:
    document = Document()
    document.add_heading("业务规范", level=1)
    document.add_paragraph(text)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _create_docx(db, session, tmp_path, monkeypatch) -> ExplorationAttachment:
    monkeypatch.setattr(settings, "uploads_dir", str(tmp_path))
    return W.create_bytes(
        db, session, "policies/rules.docx", _docx_bytes(),
        mime_type=DOCX_MIME, source="upload",
    )


def _completed(payload: dict, returncode: int = 0):
    return SimpleNamespace(
        returncode=returncode, stdout=json.dumps(payload, ensure_ascii=False), stderr="",
    )


def test_view_uses_snapshot_and_never_changes_uploaded_file(
        db, exploration_session, tmp_path, monkeypatch):
    """OfficeCLI 的 view 会规范化写回 DOCX；适配器必须把副作用困在临时副本。"""
    row = _create_docx(db, exploration_session, tmp_path, monkeypatch)
    original = Path(row.file_path).read_bytes()
    original_hash = row.sha256

    monkeypatch.setattr(O, "executable", lambda: "/fake/officecli")

    def fake_run(command, **_kwargs):
        assert command[1] == "view"
        candidate = Path(command[2])
        assert candidate != Path(row.file_path)
        candidate.write_bytes(b"officecli-normalized-the-copy")
        return _completed({"success": True, "data": {"text": "原始条款"}})

    monkeypatch.setattr(O.subprocess, "run", fake_run)
    result = O.operate(
        db, exploration_session, "view", file_id=row.id,
        view="text", start=1, max_lines=50,
    )

    assert result["result"] == {"text": "原始条款"}
    assert result["version"] == 1
    assert Path(row.file_path).read_bytes() == original
    db.refresh(row)
    assert row.sha256 == original_hash and row.version == 1


def test_edit_validates_then_atomically_updates_content_and_version(
        db, exploration_session, tmp_path, monkeypatch):
    row = _create_docx(db, exploration_session, tmp_path, monkeypatch)
    monkeypatch.setattr(O, "executable", lambda: "/fake/officecli")

    def fake_run(command, **_kwargs):
        operation = command[1]
        if operation == "validate":
            # 与真实 OfficeCLI 一致：校验通过时 data 是字符串而非对象。
            return _completed({"success": True, "data": "OpenXML validation passed"})
        assert operation == "set"
        candidate = command[2]
        document = Document(candidate)
        document.paragraphs[1].text = "修订后的条款"
        document.save(candidate)
        return _completed({"success": True, "data": {"updated": 1}})

    monkeypatch.setattr(O.subprocess, "run", fake_run)
    result = O.operate(
        db, exploration_session, "set", file_id=row.id,
        selector="/document/body/p[2]", props={"text": "修订后的条款"},
        expected_version=1,
    )

    assert result["updated"] is True
    assert result["version"] == 2
    assert result["validation"]["before"] == result["validation"]["after"] == 0
    db.refresh(row)
    assert row.version == 2 and row.status == "ready"
    assert row.source == "upload"
    assert "修订后的条款" in row.extracted_text
    assert Document(row.file_path).paragraphs[1].text == "修订后的条款"


def test_uploaded_office_remains_protected_in_later_turn_after_agent_edit(
        db, exploration_session, tmp_path, monkeypatch):
    row = _create_docx(db, exploration_session, tmp_path, monkeypatch)
    monkeypatch.setattr(O, "executable", lambda: "/fake/officecli")

    def fake_run(command, **_kwargs):
        operation = command[1]
        if operation == "validate":
            return _completed({"success": True, "data": "OpenXML validation passed"})
        assert operation == "set"
        candidate = command[2]
        document = Document(candidate)
        document.paragraphs[1].text = "用户明确授权的修订"
        document.save(candidate)
        return _completed({"success": True, "data": {"updated": 1}})

    monkeypatch.setattr(O.subprocess, "run", fake_run)
    authorized = ExplorationToolRunner(
        db, exploration_session,
        user_message="请编辑文件 rules.docx，把第二段改成已确认条款",
    )
    result = authorized.run("manage_office_document", {
        "operation": "set",
        "file_id": row.id,
        "selector": "/document/body/p[2]",
        "props": {"text": "用户明确授权的修订"},
        "expected_version": 1,
    })
    assert result["updated"] is True and result["version"] == 2
    db.refresh(row)
    assert row.source == "upload"

    monkeypatch.setattr(
        O, "operate",
        lambda *_args, **_kwargs: pytest.fail("未授权 Office 修改不应到达适配器"),
    )
    unauthorized = ExplorationToolRunner(
        db, exploration_session,
        user_message="只总结 rules.docx，不要修改或删除任何文件",
    )
    blocked_edit = unauthorized.run("manage_office_document", {
        "operation": "set",
        "file_id": row.id,
        "selector": "/document/body/p[2]",
        "props": {"text": "不应写入"},
        "expected_version": 2,
    })
    assert blocked_edit["confirmationRequired"] is True
    blocked_delete = unauthorized.run("manage_workspace_file", {
        "action": "delete",
        "file_id": row.id,
    })
    assert blocked_delete["confirmationRequired"] is True
    db.refresh(row)
    assert row.source == "upload" and row.version == 2
    assert Document(row.file_path).paragraphs[1].text == "用户明确授权的修订"


def test_edit_rejects_stale_version_without_calling_officecli(
        db, exploration_session, tmp_path, monkeypatch):
    row = _create_docx(db, exploration_session, tmp_path, monkeypatch)
    monkeypatch.setattr(
        O, "_invoke", lambda _args: pytest.fail("版本冲突时不应调用 OfficeCLI"),
    )

    with pytest.raises(HTTPException) as error:
        O.operate(
            db, exploration_session, "set", file_id=row.id,
            selector="/document/body/p[2]", props={"text": "过期写入"},
            expected_version=0,
        )
    assert error.value.status_code == 409
    assert error.value.detail["code"] == "workspace_version_conflict"


def test_validation_regression_rolls_back_original_file(
        db, exploration_session, tmp_path, monkeypatch):
    row = _create_docx(db, exploration_session, tmp_path, monkeypatch)
    original = Path(row.file_path).read_bytes()
    validate_calls = {"count": 0}
    monkeypatch.setattr(O, "executable", lambda: "/fake/officecli")

    def fake_run(command, **_kwargs):
        operation = command[1]
        if operation == "validate":
            validate_calls["count"] += 1
            if validate_calls["count"] == 1:
                return _completed({"success": True, "data": {"count": 0, "errors": []}})
            return _completed(
                {"success": False, "data": {
                    "count": 1, "errors": [{"message": "invalid relationship"}],
                }},
                returncode=1,
            )
        assert operation == "set"
        Path(command[2]).write_bytes(b"invalid-ooxml")
        return _completed({"success": True, "data": {"updated": 1}})

    monkeypatch.setattr(O.subprocess, "run", fake_run)
    with pytest.raises(HTTPException) as error:
        O.operate(
            db, exploration_session, "set", file_id=row.id,
            selector="/document/body/p[2]", props={"text": "破坏结构"},
            expected_version=1,
        )

    assert error.value.status_code == 422
    assert error.value.detail["code"] == "office_validation_regression"
    assert Path(row.file_path).read_bytes() == original
    db.refresh(row)
    assert row.version == 1


def test_external_resource_props_are_blocked_before_officecli_runs(
        db, exploration_session, tmp_path, monkeypatch):
    row = _create_docx(db, exploration_session, tmp_path, monkeypatch)
    monkeypatch.setattr(
        O, "_invoke", lambda _args: pytest.fail("危险属性不应到达 OfficeCLI"),
    )

    with pytest.raises(HTTPException, match="外部资源") as error:
        O.operate(
            db, exploration_session, "add", file_id=row.id,
            selector="/document/body", element_type="image",
            props={"imageUrl": "file:///etc/passwd"}, expected_version=1,
        )
    assert error.value.status_code == 422


def test_agent_tool_exposes_read_edit_contract_and_passes_batch_arguments(
        db, exploration_session, monkeypatch):
    operations = OFFICE_TOOL["parameters"]["properties"]["operation"]["enum"]
    assert set(operations) >= {"view", "get", "query", "validate", "replace", "batch"}
    captured = {}

    def fake_operate(_db, _session, operation, **kwargs):
        captured.update({"operation": operation, **kwargs})
        return {"operation": operation, "updated": True, "version": 3}

    monkeypatch.setattr(O, "operate", fake_operate)
    # 本用例验证参数透传；用 agent 自建文件绕开“修改用户文件需当前消息明确授权”
    # 这一独立安全门，避免不存在的 office-id 在 dispatch 前被正确拒绝。
    monkeypatch.setattr(
        W, "require_file",
        lambda _db, _session_id, _file_id: SimpleNamespace(
            id="office-id", source="agent"),
    )
    runner = ExplorationToolRunner(db, exploration_session)
    edits = [{
        "operation": "replace", "selector": "/document/body",
        "find": "旧名称", "replacement": "新名称",
    }]
    result = runner.run("manage_office_document", {
        "operation": "batch", "file_id": "office-id", "expected_version": 2,
        "edits": edits, "max_lines": 80, "cell_range": "A1:D20",
    })

    assert result["version"] == 3
    assert captured["expected_version"] == 2 and captured["edits"] == edits
    assert captured["max_lines"] == 80 and captured["cell_range"] == "A1:D20"
